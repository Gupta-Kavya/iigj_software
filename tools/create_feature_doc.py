from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\xampp\htdocs\iigj_software\IIGJ_Lab_Software_Feature_List.docx")


sections = [
    ("Core System", [
        "Fully customized gemstone laboratory management software",
        "One-time licensed custom software with no subscription dependency",
        "Multi-branch and multi-location support",
        "Location-wise user hierarchy",
        "Branch-wise data separation",
        "Users of one branch can access only their branch data",
        "Super admin can manage all branches and users",
        "Role-based access control",
        "Secure login system",
        "Signup removed for controlled access",
        "IP-based access restriction for security",
        "Branch-wise IP allowlist management",
        "Compact customized login screen with GJEPC branding",
    ]),
    ("Branch And Location Management", [
        "Add and manage multiple branch locations",
        "Store branch name, address, phone, email, website, CIN, and GST",
        "Branch details automatically appear on agreement print",
        "Branch-wise reports, agreements, customers, images, and dashboard data",
        "Branch-wise image folders",
        "Collection center support",
        "Collection center code/letter used in report/reference number",
        "Receiving branch and collection center can be handled separately",
    ]),
    ("User Management", [
        "Create branch users from super admin",
        "Assign every user to a branch/location",
        "Manage user role: User or Super Admin",
        "Activate or deactivate users",
        "Reset user password",
        "View full user profile",
        "User-wise activity insights",
        "User-wise agreements, reports, customers, images, and financial summary",
        "Last login tracking",
        "User status tracking",
    ]),
    ("Agreement Module", [
        "Create stone agreement before certificate generation",
        "Agreement-cum-acknowledgement receipt format",
        "Exact IIGJ-style agreement print layout",
        "Customer copy format",
        "Agreement number generation",
        "Date and time capture",
        "Customer name selection from customer master",
        "Auto-fill customer details after selection",
        "Add new customer directly from agreement form",
        "Customer mobile, address, GST, email, and ID details",
        "Urgent / Regular selection",
        "Member / Non-member selection",
        "MOU / CDC selection",
        "Silver / Gold / Platinum card support",
        "Membership discount support",
        "Dynamic stone rows",
        "Add multiple stone entries",
        "Cancel agreement row instead of deleting it",
        "Cancelled rows retained in records",
        "Cancelled row amount deducted from total",
        "Cancelled amount added to refund",
        "Cancellation reason popup",
        "WhatsApp notification on row cancellation",
        "Agreement status management",
        "Status values such as In Process, Delivered, and Cancelled",
        "Delivery date and time auto-update on delivered status",
        "WhatsApp notification on status update",
        "Agreement edit option",
        "Agreement print option",
        "Agreement label generation option",
        "Save confirmation popup before final agreement save",
        "Important agreement details shown before saving",
        "After-save option to generate agreement, generate labels, or exit",
        "Agreement total PCS calculation",
        "Total testing charges calculation",
        "Discount calculation",
        "Due amount calculation",
        "Refund amount calculation",
        "Amount in words shown on print",
        "Payment mode handling: cash, cheque, NEFT/UPI, card, and TDS",
        "Manual signature support",
        "E-signature pad support",
        "Depositor signature shown on agreement",
        "IIGJ signature section",
        "Print-ready A4 agreement layout",
    ]),
    ("Agreement Row And Stone Details", [
        "Reference number generation from agreement number, collection center/location letter, and certificate number",
        "Category dropdown from rate list master",
        "Dynamic rate fetching",
        "Gross weight with unit",
        "Stone weight with unit",
        "Diamond weight fixed in carat",
        "Color suggestion from color master",
        "A4 / ATM Card selection",
        "Top-up checkbox",
        "PCS entry",
        "Automatic total PCS",
        "Automatic total amount",
        "Row-wise discount",
        "Row-wise amount calculation",
        "Cancel row option in edit mode",
        "Validation before adding rows",
        "Customer selection required before stone rows",
        "Rate condition validation before save",
    ]),
    ("Rate List And Charge Calculation", [
        "Rate list master import/use",
        "Category-wise rate selection",
        "Rate code based calculation",
        "Amount auto-calculation",
        "Packet lot minimum PCS validation",
        "Bead length based calculation",
        "Minimum bead length condition",
        "Diamond packet screening logic",
        "Diamond grading minimum charge logic",
        "Top-up calculation logic",
        "Branch-specific rate condition support",
        "Dashboard-based rate condition master",
        "Super-admin-only condition management",
        "Dynamic branch/location dropdown in conditions",
        "Priority-based condition execution",
        "Value1 / Value2 based rule setup",
        "Future condition addition without code changes",
    ]),
    ("Customer Master", [
        "Customer master created from Excel",
        "Customer search/suggestion on agreement form",
        "Auto-fill customer details",
        "Customer branch-wise access",
        "Customer mobile, address, GST, email, and ID details",
        "Membership details",
        "MOU/CDC card type",
        "Discount percentage mapping",
        "Add customer from agreement screen",
    ]),
    ("Certificate Booking Flow", [
        "Agreement automatically books certificates/reports",
        "Booked certificates stored separately",
        "Report/reference number linked with agreement",
        "Certificate number continues from last generated number",
        "Prevents incorrect certificate entry",
        "Feeding forms check agreement number and certificate number",
        "Existing certificate can be loaded and edited",
        "Prevents wrong report type feeding in another form",
        "Branch-wise certificate visibility",
    ]),
    ("Color Stone Feeding", [
        "Compact color stone feeding form",
        "Agreement number and certificate number required before data entry",
        "Fields disabled until valid booking is entered",
        "Auto-fetch booked agreement data",
        "Edit already generated certificate",
        "Dynamic report type dropdown",
        "Report type controls certificate format",
        "Stone image fetch/capture support",
        "Camera/image upload support",
        "Item description, gross weight, color, shape/cut, and tested PCS",
        "Stone weight 1 to 5 with unit handling",
        "Measurements 1 to 5",
        "Length tested, RI, SG, and optic character",
        "Species / Variety mode and Others mode",
        "Group/species and variety/others handling",
        "Origin, treatment description, and general comments master",
        "Test carried out tick fields",
        "Save/update logic",
        "After-save confirmation to generate report or exit",
    ]),
    ("Pearl Feeding", [
        "Pearl report feeding based on color stone flow",
        "Type stored as Pearl",
        "Agreement/certificate validation",
        "Booking fetch logic",
        "Save/update logic",
        "Report generation flow",
        "Image handling similar to color stone",
    ]),
    ("Diamond Jewellery Feeding", [
        "Diamond jewellery feeding form",
        "Agreement number and certificate number validation",
        "Booking fetch logic",
        "Existing certificate edit logic",
        "Prevents loading wrong report type",
        "Dynamic report type dropdown",
        "Item details, metal type/gold purity, and gross weight",
        "Diamond weight, shape/cut, diamond color, clarity, and finish",
        "Stone name master connection",
        "Color stone details 1 to 7",
        "Color 1 to 7, weight 1 to 7, and stone name 1 to 7",
        "Comment field",
        "Stone image fetch/capture",
        "Save/update logic",
        "Generate report confirmation after save",
    ]),
    ("Diamond Screening Feeding", [
        "Diamond screening feeding form",
        "Agreement number and certificate number validation",
        "Booking fetch logic",
        "Existing report edit support",
        "Dynamic report type dropdown",
        "Shape and cut",
        "Total weight",
        "Total PCS stored properly",
        "Natural diamond weight and PCS",
        "Synthetic diamond weight and PCS",
        "Referral weight and PCS",
        "Non-diamond weight and PCS",
        "Stone image fetch/capture",
        "Save/update logic",
        "Generate certificate or labels after save",
        "Diamond screening label format",
    ]),
    ("Diamond Grading Feeding", [
        "Diamond grading feeding form",
        "Agreement number and certificate number validation",
        "Dynamic report type dropdown",
        "Natural/synthetic handled through report type",
        "Item description, shape and cut, measurement, and carat weight",
        "Cut, symmetry, polish, table size, crown height, and pavilion depth",
        "Girdle thickness, culet, color, fluorescence, and clarity",
        "General comments",
        "Stone image, proportion image, and clarity image",
        "Separate image folders",
        "Auto-fetch images by certificate number",
        "Symbol master support",
        "Select up to 3 symbols",
        "Symbols stored in reusable structured format",
        "Symbol images fetched from folder",
        "Test carried out tick fields",
        "Save/update logic",
        "Report generation flow",
    ]),
    ("Report Type Master", [
        "Dynamic report type creation",
        "Report types can be created for color stone, pearl, diamond jewellery, diamond screening, and diamond grading",
        "Report type decides report format",
        "Report type decides A4 / ATM / Postcard / PVC Card style",
        "Feeding forms show only relevant report types",
        "Report builder automatically receives report type formats",
    ]),
    ("Report Builder", [
        "Drag-and-drop report builder",
        "A4, ATM card, postcard, and PVC card report layouts",
        "Postcard size support: 10 cm x 15 cm",
        "Field, label, value, tick mark, QR code, and image positioning",
        "Dynamic tick marks based on feeding checkboxes",
        "Stone, proportion, clarity, symbol, and additional image placement",
        "Additional image upload and delete option",
        "Additional text placement",
        "All database fields available in builder",
        "Color, comments, and treatment description fields available",
        "No hardcoded limited fields",
        "Label font and value font control",
        "Individual label and value font color",
        "Font size, font weight, and bold option",
        "Font file upload and custom TTF font support",
        "Arial Nova Cond Light font support",
        "Hide/show colon between label and value",
        "Label width control",
        "Value wrapping support",
        "Display condition support for labels/text",
        "AND / OR condition support",
        "Report builder settings backup export",
        "Report builder settings import/restore",
        "Builder does not refresh page after saving",
        "Improved positioning accuracy between builder and final output",
        "Scrollable preview/editor layout",
        "Side pane editing support",
    ]),
    ("Report Generation", [
        "Generate reports from feeding forms",
        "Generate reports separately from dashboard",
        "Auto-generate by report type",
        "PVC card, A4, ATM card, and postcard report generation",
        "Certificate preview",
        "Print-ready output",
        "QR code support",
        "Report images support",
        "Branch-wise report format usage",
        "Location-wise report builder settings",
        "Same branch users share same report formats",
        "Users from other branches cannot see other branch reports",
    ]),
    ("Image Management", [
        "Stone image management",
        "Symbol image management",
        "Clarity image management",
        "Proportion image management",
        "Branch-wise image folders",
        "No unnecessary user-wise folders",
        "Images stored under branch/location",
        "Auto-fetch image by certificate number",
        "Camera capture support",
        "Image upload support",
        "View uploaded images",
        "Use images in report builder",
    ]),
    ("Label Printing", [
        "Agreement label generation",
        "Label design for label printer",
        "Compact label layout",
        "75 mm x 38 mm label PDF support",
        "Fixed-size label PDF generation",
        "Rec date and time",
        "Delivery date and time",
        "Reference number and serial number",
        "PCS for testing",
        "Stone weight",
        "Category",
        "Report type",
        "Document number",
        "Diamond screening label generation",
        "Proper label font size",
        "Borderless neat label design",
        "Print-ready label output",
    ]),
    ("WhatsApp Integration", [
        "WaAPI integration",
        "Super admin can set instance ID and API key",
        "Send agreement on WhatsApp",
        "Send to multiple mobile numbers separated by comma",
        "Send status update WhatsApp messages",
        "Send cancellation WhatsApp message with reason",
        "Formatted WhatsApp message",
        "Customer name and agreement details in message",
        "Preview link disabled",
        "Toast notification after sending",
        "No default browser alert",
        "WhatsApp document/media sending support",
    ]),
    ("Dashboard", [
        "Compact dashboard layout",
        "Detailed insights with graphics",
        "Reports this month",
        "Reports entered today",
        "Monthly volume",
        "Report composition",
        "Daily report trend",
        "Agreement insights",
        "Customer insights",
        "User activity insights",
        "Branch-wise overview",
        "Financial summary",
        "Testing charges, due amount, and refund amount",
        "Report type composition",
        "Useful operational metrics",
    ]),
    ("Super Admin Panel", [
        "User account management",
        "Branch/location management",
        "Allowed IP management",
        "WhatsApp API settings",
        "Report type management",
        "Rate condition management",
        "Branch-wise details",
        "User details and user activity",
        "Agreement, report, customer, and image statistics",
        "Due/refund/testing totals",
        "Recent reports",
        "Recent agreements",
        "Recent customers",
    ]),
    ("Security", [
        "Login authentication",
        "Super Admin and User roles",
        "IP allowlist",
        "Branch-wise access restriction",
        "No public signup",
        "Data isolation by branch",
        "Controlled user creation",
        "Password reset by admin",
        "Session-based access",
        "Unauthorized access blocking",
    ]),
    ("Data And Reporting Benefits", [
        "Separate agreement table",
        "Separate agreement child/row table",
        "Certificate data stored for reporting",
        "Stone rows stored in report-friendly format",
        "Cancelled rows retained for audit",
        "Reports can be generated later",
        "Branch-wise, customer-wise, stone-wise, category-wise, PCS-wise, amount-wise, and status-wise reports possible",
        "Future MIS reports possible",
    ]),
    ("Operational Benefits", [
        "Reduces manual agreement work",
        "Reduces certificate formatting mistakes",
        "Standardized report output",
        "Faster lab feeding process",
        "Better branch control",
        "Better customer communication",
        "Better record keeping",
        "Easy future report format changes",
        "Easy future rate condition changes",
        "Proper audit trail for cancelled rows",
        "Professional printed agreements, reports, and labels",
        "Designed for gemstone lab workflow, not generic billing software",
    ]),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="DADCE0"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_run(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(item)
        style_run(r, size=9.4, color="1F2937")


def add_section(doc, title, items):
    h = doc.add_heading(title, level=1)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    h.runs[0].font.color.rgb = RGBColor(31, 77, 120)
    h.runs[0].font.size = Pt(14)
    add_bullets(doc, items)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for style_name in ("Heading 1", "Heading 2"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("IIGJ Gem Testing Lab Software")
    style_run(r, size=23, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("Complete Feature Catalogue For Client Proposal")
    style_run(r, size=11, bold=True, color="5B677A")

    pitch = doc.add_table(rows=1, cols=1)
    pitch.alignment = WD_TABLE_ALIGNMENT.CENTER
    pitch.autofit = True
    cell = pitch.cell(0, 0)
    set_cell_shading(cell, "EEF4FB")
    set_cell_border(cell, "B8C7D9")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "This is not just certificate printing software. It is a complete gemstone laboratory workflow system covering customer intake, agreement creation, sample booking, rate calculation, certificate feeding, report designing, image handling, label printing, WhatsApp communication, branch-wise access control, and management reporting."
    )
    style_run(r, size=10.5, bold=True, color="0B2545")

    doc.add_paragraph()

    summary = doc.add_table(rows=1, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.autofit = False
    summary.columns[0].width = Inches(2.2)
    summary.columns[1].width = Inches(4.55)
    for c in summary.row_cells(0):
        set_cell_border(c, "DADCE0")
    summary.cell(0, 0).text = "Software Type"
    summary.cell(0, 1).text = "Customized gemstone laboratory management and certificate workflow system"
    row = summary.add_row().cells
    row[0].text = "Primary Users"
    row[1].text = "Testing labs, branch users, report feeding users, super admins, and management"
    row = summary.add_row().cells
    row[0].text = "Core Value"
    row[1].text = "End-to-end control from agreement to certificate, label, WhatsApp communication, reporting, and branch-wise MIS"
    for row in summary.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_border(cell, "DADCE0")
            if idx == 0:
                set_cell_shading(cell, "F2F4F7")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    style_run(run, size=9.2, bold=(idx == 0), color="111827")

    for title, items in sections:
        add_section(doc, title, items)

    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(10)
    closing.paragraph_format.space_after = Pt(0)
    r = closing.add_run("Positioning Statement: ")
    style_run(r, size=10.5, bold=True, color="0B2545")
    r = closing.add_run(
        "The system is built specifically for gemstone lab operations and can be extended with future report formats, rate rules, branches, images, and MIS requirements."
    )
    style_run(r, size=10.5, color="111827")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "IIGJ Gem Testing Lab Software - Feature Catalogue"
    for run in footer.runs:
        style_run(run, size=8.5, color="6B7280")

    doc.save(OUT)


if __name__ == "__main__":
    build()
